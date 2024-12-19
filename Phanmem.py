from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QMessageBox, QFileDialog
from PyQt5.QtCore import QThread, pyqtSignal
import pandas as pd
import sys
import os
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog
from PyQt5.QtGui import QTextDocument, QTextCursor
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import subprocess
import openai
import yaml
from docx import Document
import re
import shutil
import dotenv


dotenv.load_dotenv()  # Load environment variables from .env file

openai.api_key = os.getenv("OPENAI_API_KEY")

def AI_Func(df):
    def load_conditions(file_path):
        """Loads conditions from a .docx file.
        Assumes YAML content is in a code block (within double backticks).
        """
        doc = Document(file_path)
        yaml_content = ""

        for paragraph in doc.paragraphs:
            if paragraph.text.startswith("```") and paragraph.text.endswith("```"):
                yaml_content = paragraph.text[3:-3]
                break

        if yaml_content:
            conditions = yaml.safe_load(yaml_content)
            return conditions
        else:
            raise ValueError("No YAML content found within code block in docx file.")


    def save_conditions(file_path, conditions):
        """Saves conditions to a .docx file.
        Saves YAML content within a code block (in double backticks).
        """
        doc = Document()
        doc.add_paragraph(f"```\n{yaml.dump(conditions)}\n```")
        doc.save(file_path)


    def check_sudden_drop(df, load_name, threshold=500):
        """Check if a load has a sudden drop in the last month."""
        thang_9_load = df[df["Tên"] == load_name]["Tháng 9"].values[0]
        thang_8_load = df[df["Tên"] == load_name]["Tháng 8"].values[0]
        return (thang_8_load - thang_9_load) > threshold


    def get_llm_choice(potential_loads_df, highest_phase, lowest_phase, conditions_text):
        """Asks the LLM to choose the best load to move."""

       
        loads_info = ""
        for index, row in potential_loads_df.iterrows():
            loads_info += f"  - {row['Tên']}: Tải Tháng 9 = {row['Tháng 9']}, Giảm đột ngột = {row['Giảm đột ngột']}\n"

        prompt = (
            f"Bạn là một chuyên gia trong việc cân bằng tải lưới điện.\n"
            f"Giúp tôi chọn tải tốt nhất để di chuyển từ pha cao nhất ({highest_phase}) "
            f"đến pha thấp nhất ({lowest_phase}) để tối ưu hóa sự cân bằng, xem xét các điều kiện sau:\n\n"
            f"{conditions_text}\n\n"
            f"Các tải tiềm năng (được sắp xếp theo thứ tự ưu tiên)::\n"
            f"{loads_info}\n"
            f"Chọn **tên** của tải tốt nhất để di chuyển, đảm bảo nó đáp ứng tất cả các điều kiện. "
            f"Nếu không tìm thấy tải phù hợp, hãy phản hồi bằng 'Không có'."
        )

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000
        )

        choice = response['choices'][0]['message']['content'].strip()
    
       
        match = re.search(r"(Load_\d+)", choice)  
        if match:
            choice = match.group(1)
        else:
            choice = "None"  
    
        return choice


    def balance_phases(df, conditions_text, max_iterations=15, max_moves_per_load=3):
        """Balances phases using rules and LLM consultation."""

        df["Pha di chuyển"] = "" 
        df["Pha hiện tại"] = df["Pha"].copy()  
        iteration = 0

        while iteration < max_iterations:
            phase_sums = df.groupby("Pha")["Tháng 9"].sum()
            highest_phase = phase_sums.idxmax()
            lowest_phase = phase_sums.idxmin()

            print(f"\nLần lặp {iteration + 1}:")
            print(f"Tổng pha hiện tại: {phase_sums.to_dict()}")

            if (phase_sums.max() - phase_sums.min()) <= 200:
                print("Các pha đã được cân bằng. Thoát.")
                break

            target_value = (phase_sums[highest_phase] - phase_sums[lowest_phase]) / 2

            potential_loads = df[df["Pha"] == highest_phase].copy()
            potential_loads["Khoảng cách"] = (potential_loads["Tháng 9"] - target_value).abs()
            potential_loads = potential_loads[potential_loads["Pha di chuyển"].str.count(",") < max_moves_per_load] 

          
            potential_loads = potential_loads.sort_values(
                by=["Khoảng cách", "Tháng 9", "Giảm đột ngột"], ascending=[True, True, False]
            )

            
            llm_choice = get_llm_choice(potential_loads, highest_phase, lowest_phase, conditions_text)

            print(f"LLM đã chọn di chuyển tải: {llm_choice}")

            if llm_choice != "None" and llm_choice in df["Tên"].values:
                df.loc[df["Tên"] == llm_choice, "Pha"] = lowest_phase
                df.loc[df["Tên"] == llm_choice, "Pha di chuyển"] += (
                    f"{highest_phase} sang {lowest_phase}" 
                    if df.loc[df["Tên"] == llm_choice, "Pha di chuyển"].iloc[0] == "" 
                    else f", {highest_phase} sang {lowest_phase}"
                )
            elif llm_choice == "None":
                print("LLM không tìm thấy tải phù hợp để di chuyển trong lần lặp này.")
            else:
                print(f"LLM đề xuất tải không hợp lệ: {llm_choice}")

            iteration += 1

        if iteration == max_iterations:
            print("Đạt đến số lần lặp tối đa. Thoát.")

       
        df['Pha đề xuất'] = df['Pha'].copy()  

        df_balanced = df[['Tên', 'Khách hàng', 'Mã KH', 'Số công tơ', 'Sổ ghi số', 'Tháng 6', 'Tháng 7', 'Tháng 8', 'Tháng 9', 'Pha hiện tại', 'Pha di chuyển', 'Pha đề xuất']] 
        return df_balanced


    
    df = pd.read_excel("table1.xlsx")
    if 'Pha hiện tại' not in df.columns: 
        df["Pha hiện tại"] = df["Pha"].copy()
    df['Tháng 6'] = pd.to_numeric(df['Tháng 6'], errors='coerce')
    df['Tháng 7'] = pd.to_numeric(df['Tháng 7'], errors='coerce')
    df['Tháng 8'] = pd.to_numeric(df['Tháng 8'], errors='coerce')
    df['Tháng 9'] = pd.to_numeric(df['Tháng 9'], errors='coerce')

    df["Giảm đột ngột"] = df["Tên"].apply(lambda x: check_sudden_drop(df, x))

    
    conditions_text = """
    Vấn đề: Các pha không cân bằng.
    Mục tiêu: Cân bằng tải trên các pha để đảm bảo ổn định. 
    Ưu tiên:
    1. Gần nhất với giá trị mục tiêu.
    2. Không nằm trong top 3 tải cao nhất.
    3. Không giảm đột ngột trong Tháng 8.
    """

    df_balanced = balance_phases(df.copy(), conditions_text)

    print("\n\nDữ liệu cân bằng cuối cùng:\n", df_balanced.to_string())
    print("\nTổng pha cân bằng:\n", df_balanced.groupby("Pha đề xuất")["Tháng 9"].sum())  

    return df_balanced


class Ui_Form_PickTram(object):
    def setupUi_PickTram(self, Form):
        Form.setObjectName("Form")
        Form.resize(660, 500)
        self.chontram = QtWidgets.QFrame(Form)
        self.chontram.setGeometry(QtCore.QRect(0, 0, 661, 501))
        self.chontram.setStyleSheet("background-color: rgb(0, 0, 139);")
        self.chontram.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.chontram.setFrameShadow(QtWidgets.QFrame.Raised)
        self.chontram.setObjectName("chontram")
        self.name = QtWidgets.QLabel(self.chontram)
        self.name.setGeometry(QtCore.QRect(140, 40, 361, 111))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(40)
        font.setBold(True)
        font.setWeight(75)
        self.name.setFont(font)
        self.name.setStyleSheet("color: rgb(255, 255, 255);")
        self.name.setObjectName("name")
        self.LeNgocHan = QtWidgets.QPushButton(self.chontram)
        self.LeNgocHan.setGeometry(QtCore.QRect(180, 200, 281, 61))
        font = QtGui.QFont()
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.LeNgocHan.setFont(font)
        self.LeNgocHan.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 139);")
        self.LeNgocHan.setObjectName("LeNgocHan")												   
        self.ConditionText = QtWidgets.QPushButton(self.chontram)
        self.ConditionText.setGeometry(QtCore.QRect(180, 330, 281, 71))
        font = QtGui.QFont()
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.ConditionText.setFont(font)
        self.ConditionText.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 139);")
        self.ConditionText.setObjectName("ConditionText")

        self.retranslateUi_PickTram(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)
        self.LeNgocHan.clicked.connect(self.load_data_doi_can)
        self.ConditionText.clicked.connect(self.load_data_condition)
        
    def load_data_doi_can(self):
        self.load_data('table1')

    def load_data_condition(self):
        if os.path.exists("condition.docx"):
            self.open_docx("condition.docx")
        else:
            options = QFileDialog.Options()
            options |= QFileDialog.ReadOnly
            fileName, _ = QFileDialog.getOpenFileName(None, "Chọn tệp điều kiện", "", 
                                                      "Word Files (*.docx)", options=options)
            if fileName:
                shutil.copyfile(fileName, "condition.docx")
                self.open_docx("condition.docx")

    def load_data(self, table_name, file_filter="Excel Files (*.xlsx)"):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        fileName, _ = QFileDialog.getOpenFileName(None, "QFileDialog.getOpenFileName()", "", file_filter, options=options)
        if fileName:
            if file_filter == "Excel Files (*.xlsx)":
                df = pd.read_excel(fileName)
                df.to_excel(f'{table_name}.xlsx', index=False)

                msg = QMessageBox()
                msg.setIcon(QMessageBox.Information)
                msg.setText("Dữ liệu đã được cập nhật thành công!")
                msg.setWindowTitle("Thành công")
                msg.exec_()
            else:
                
                self.open_docx(fileName)

    def open_docx(self, file_path):
        try:
            subprocess.Popen([file_path], shell=True)
        except Exception as e:
            print(f"Lỗi khi mở tệp: {e}")

    def retranslateUi_PickTram(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Form"))
        self.name.setText(_translate("Form", "Chọn trạm:"))
        self.LeNgocHan.setText(_translate("Form", "Lê Ngọc Hân"))
        self.ConditionText.setText(_translate("Form", "Điều kiện xác định"))
        
class Ui_Form_ResultFinal(object):
    def setupUi_ResultFinal(self, Form, df_balanced): 
        Form.setObjectName("Form")
        Form.resize(1920, 1080)
        self.frame = QtWidgets.QFrame(Form)
        self.frame.setGeometry(QtCore.QRect(0, 0, 1920, 1080))
        self.frame.setStyleSheet("background-color: rgb(0, 0, 139);")	
        self.frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame.setObjectName("frame")
        self.tieudecuabangdulieu = QtWidgets.QFrame(self.frame)
        self.tieudecuabangdulieu.setGeometry(QtCore.QRect(100, 99, 761, 41))
        self.tieudecuabangdulieu.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.tieudecuabangdulieu.setFrameShadow(QtWidgets.QFrame.Raised)
        self.tieudecuabangdulieu.setObjectName("tieudecuabangdulieu")												 
        self.PICK_AGAIN = QtWidgets.QPushButton(self.frame)
        self.PICK_AGAIN.setGeometry(QtCore.QRect(770, 950, 321, 41))
        font = QtGui.QFont()
        font.setPointSize(16)
        font.setBold(True)
        font.setWeight(75)
        self.PICK_AGAIN.setFont(font)
        self.PICK_AGAIN.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 139);")
        self.PICK_AGAIN.setObjectName("PICK_AGAIN")
        self.PRINT = QtWidgets.QPushButton(self.frame)
        self.PRINT.setGeometry(QtCore.QRect(770, 890, 321, 41))
        font = QtGui.QFont()
        font.setPointSize(16)
        font.setBold(True)
        font.setWeight(75)
        self.PRINT.setFont(font)
        self.PRINT.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 139);")
        self.PRINT.setObjectName("PRINT")														 
        self.logo = QtWidgets.QLabel(self.frame)
        self.logo.setGeometry(QtCore.QRect(1140, 0, 141, 131))
        self.logo.setText("")
        self.logo.setPixmap(QtGui.QPixmap("OneDrive/Tài liệu/D14TDHHTD2/Đồ án tốt nghiệp/APP/.designer/.designer/lapso/Downloads/snapedit_1701002446534.png"))
        self.logo.setObjectName("logo")
        self.RESULT_TABLE = QtWidgets.QTableWidget(self.frame)  
        self.RESULT_TABLE.setGeometry(QtCore.QRect(320, 290, 971, 581))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.RESULT_TABLE.setFont(font)
        self.RESULT_TABLE.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 0);")
        self.RESULT_TABLE.setObjectName("RESULT_TABLE")
        self.RESULT_TABLE.setColumnCount(0)
        self.RESULT_TABLE.setRowCount(0)
        self.RESULT_TABLE.setRowCount(df_balanced.shape[0])
        self.RESULT_TABLE.setColumnCount(df_balanced.shape[1])
        self.RESULT_TABLE.setHorizontalHeaderLabels(df_balanced.columns.tolist())
        
        for i, row in df_balanced.iterrows():
            for j, value in enumerate(row):
                self.RESULT_TABLE.setItem(i, j, QtWidgets.QTableWidgetItem(str(value)))												 
        self.tentieude_2 = QtWidgets.QLabel(self.frame)
        self.tentieude_2.setGeometry(QtCore.QRect(1450, 220, 331, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.tentieude_2.setFont(font)
        self.tentieude_2.setStyleSheet("color: rgb(255, 255, 255);")
        self.tentieude_2.setAlignment(QtCore.Qt.AlignCenter)
        self.tentieude_2.setObjectName("tentieude_2")
		
        self.RESULT_TABLE.update()
        self.ResultFinalForm = Form
        self.df_balanced = df_balanced  
        self.PRINT.clicked.connect(self.save_as_excel)
        self.PICK_AGAIN.clicked.connect(self.replace_df_with_table3)  
        self.textEdit = QtWidgets.QTextEdit(self.frame)
        self.textEdit.setGeometry(QtCore.QRect(1340, 290, 541, 581))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(14)
        self.textEdit.setFont(font)
        self.textEdit.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 0);")
        self.textEdit.setObjectName("textEdit")
        sys.stdout.write = lambda message: self.textEdit.append(str(message))
        sys.stderr.write = lambda message: self.textEdit.append(str(message))													   
        self.PRINTER = QtWidgets.QPushButton(self.frame)
        self.PRINTER.setGeometry(QtCore.QRect(1680, 890, 201, 41))
        self.dudoan = QtWidgets.QPushButton(self.frame)
        self.dudoan.setGeometry(QtCore.QRect(1440, 890, 201, 41))
        font = QtGui.QFont()
        font.setPointSize(16)
        font.setBold(True)
        font.setWeight(75)
        self.PRINTER.setFont(font)
        self.PRINTER.setStyleSheet("background-color: rgb(0, 0, 255);\n"         
        "color: rgb(255, 255, 255);")
        self.dudoan.setFont(font)
        self.dudoan.setStyleSheet("background-color: rgb(0, 0, 255);\n"         
        "color: rgb(255, 255, 255);")
        self.PRINTER.setObjectName("PRINTER")
        self.dudoan.setObjectName("dudoan")
        self.tenappviettat = QtWidgets.QLabel(self.frame)
        self.tenappviettat.setGeometry(QtCore.QRect(0, 0, 1920, 130))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(48)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.tenappviettat.setFont(font)
        self.tenappviettat.setStyleSheet("color: rgb(255, 255, 255);")
        self.tenappviettat.setTextFormat(QtCore.Qt.AutoText)
        self.tenappviettat.setAlignment(QtCore.Qt.AlignCenter)
        self.tenappviettat.setObjectName("tenappviettat")
        self.tenappviettat_2 = QtWidgets.QLabel(self.frame)
        self.tenappviettat_2.setGeometry(QtCore.QRect(0, 110, 1920, 61))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.tenappviettat_2.setFont(font)
        self.tenappviettat_2.setStyleSheet("color: rgb(255, 255, 255);")
        self.tenappviettat_2.setTextFormat(QtCore.Qt.AutoText)
        self.tenappviettat_2.setAlignment(QtCore.Qt.AlignCenter)
        self.tenappviettat_2.setObjectName("tenappviettat_2")
        self.tentieude = QtWidgets.QLabel(self.frame)
        self.tentieude.setGeometry(QtCore.QRect(400, 220, 791, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.tentieude.setFont(font)
        self.tentieude.setStyleSheet("color: rgb(255, 255, 255);")
        self.tentieude.setAlignment(QtCore.Qt.AlignCenter)
        self.tentieude.setObjectName("tentieude")

        self.retranslateUi_ResultFinal(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)
        self.PRINTER.clicked.connect(self.print_file)

        
    def replace_df_with_table3(self):  
       df_new = pd.read_excel('table3.xlsx')
       df_new.to_excel('table1.xlsx', index=False)
       df_new.to_excel('table2.xlsx', index=False)

       self.df_balanced = df_new  
       self.RESULT_TABLE.setRowCount(self.df_balanced.shape[0])  
       self.RESULT_TABLE.setColumnCount(self.df_balanced.shape[1])  
       self.RESULT_TABLE.setHorizontalHeaderLabels(self.df_balanced.columns.tolist())  

       for i, row in self.df_balanced.iterrows(): 
           for j, value in enumerate(row):
               self.RESULT_TABLE.setItem(i, j, QtWidgets.QTableWidgetItem(str(value)))

       self.RESULT_TABLE.update()
       QtWidgets.QApplication.processEvents()  
       self.ResultFinalForm.close()


        
  
    def save_as_excel(self):
          filename, _ = QFileDialog.getSaveFileName(QtWidgets.QWidget(), "Lưu tệp Excel", "", "Excel Files (*.xlsx)")
          if filename:
            self.df_balanced.to_excel(filename, index=False)  
            
    def print_file(self):
        printer = QPrinter(QPrinter.HighResolution)
        dialog = QPrintDialog(printer, self.ResultFinalForm)
    
        if dialog.exec_() == QPrintDialog.Accepted:
            text_document = QTextDocument(self.textEdit.toPlainText())
            cursor = QTextCursor(text_document)
            cursor.movePosition(QTextCursor.End)
            cursor.insertBlock()  
            cursor.insertText("\nKý xác nhận: _______________________________")  
            text_document.print_(printer)



    def retranslateUi_ResultFinal(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Form"))
        self.PICK_AGAIN.setText(_translate("Form", "Chọn lại"))
        self.PRINT.setText(_translate("Form", "Lưu excel"))
        self.tentieude_2.setText(_translate("Form", "Pha đã di chuyển"))
        self.PRINTER.setText(_translate("Form", "In tải chuyển"))
        self.tenappviettat.setText(_translate("Form", "<html><head/><body><p>Phần mềm cân bằng pha</p></body></html>"))
        self.tenappviettat_2.setText(_translate("Form", "<html><head/><body><p>Tưởng Gia Huy-Trường đại học điện lực</p></body></html>"))
        self.tentieude.setText(_translate("Form", "Phương án cân bằng pha đề xuất"))
        self.dudoan.setText(_translate("Form", "Dự Đoán Từ AI"))
        
        self.textEdit.textChanged.connect(self.append_llm_explanation) 
        self.dudoan.clicked.connect(self.predict_next_month)
        

    def append_llm_explanation(self):
        """Appends the LLM's explanation if it's not already present."""
        if "Model AI respond:" not in self.textEdit.toPlainText():  
            return 

       
        self.textEdit.textChanged.disconnect(self.append_llm_explanation) 

        llm_explanation = self.get_llm_explanation()
        self.textEdit.append(llm_explanation)
        
    def get_llm_explanation(self):
        """Asks the LLM to explain what it did to balance the phases."""

      
        changed_loads = self.df_balanced[self.df_balanced['Pha hiện tại'] != self.df_balanced['Pha đề xuất']]
        changed_loads_info = ""
        for index, row in changed_loads.iterrows():
            changed_loads_info += f"  - {row['Tên']}: từ {row['Pha hiện tại']} sang {row['Pha đề xuất']}\n"

        prompt = (
            f"Bạn là một chuyên gia trong việc cân bằng tải lưới điện. Trước đây, bạn đã được giao nhiệm vụ cân bằng. "
            f"một hệ thống ba pha. Dựa trên những thay đổi mà bạn đã đề xuất (được liệt kê bên dưới), hãy giải thích chi tiết "
            f"những gì bạn đã làm để cân bằng hệ thống và tại sao bạn lại đưa ra những lựa chọn cụ thể đó.\n\n"
            f"Những thay đổi đã thực hiện:\n"
            f"{changed_loads_info}\n\n"
            f"Giải thích:"
        )

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000
        )

        explanation = response['choices'][0]['message']['content'].strip()
        return explanation
    
    def predict_next_month(self):
        """Sends a prompt to the LLM to predict next month's situation and appends the response."""
    
        prompt = "Hãy thử dự đoán điều gì sẽ xảy ra trong tháng tới liên quan đến việc phân bố tải, và có thiên tai gì ảnh hưởng nặng nề không trong hệ thống 3 pha này không"
    
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=500 
        )
    
        prediction = response['choices'][0]['message']['content'].strip()
        self.textEdit.append(f"Dự đoán từ AI là:\n{prediction}")
        
class LongOperationThread(QThread):
   finished = pyqtSignal(object)

   def __init__(self, selected_text):
       QThread.__init__(self)
       self.selected_text = selected_text

   def run(self):
       df = None
       print("đang được phát triển")
       if self.selected_text == "Lê Ngọc Hân ":
           df = pd.read_excel('table1.xlsx')
           df = AI_Func(df)
       elif self.selected_text == "Điều kiện xác định":
           df = pd.read_excel('table2.xlsx')
           df = AI_Func(df)

       self.finished.emit(df)


class Ui_Form_YesOrNo(object):
    def setupUi_YesOrNo(self, Form, df):
        Form.setObjectName("Form")
        Form.resize(330, 250)
        self.chonkieunhaplieu = QtWidgets.QFrame(Form)
        self.chonkieunhaplieu.setGeometry(QtCore.QRect(0, 0, 330, 250))
        self.chonkieunhaplieu.setStyleSheet("background-color: rgb(85, 85, 255);")
        self.chonkieunhaplieu.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.chonkieunhaplieu.setFrameShadow(QtWidgets.QFrame.Raised)
        self.chonkieunhaplieu.setObjectName("chonkieunhaplieu")
        self.chonkieunhapdulieu = QtWidgets.QLabel(self.chonkieunhaplieu)
        self.chonkieunhapdulieu.setGeometry(QtCore.QRect(50, 10, 251, 41))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(15)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu.setFont(font)
        self.chonkieunhapdulieu.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu.setObjectName("chonkieunhapdulieu")
        self.NOT = QtWidgets.QPushButton(self.chonkieunhaplieu)
        self.NOT.setGeometry(QtCore.QRect(90, 60, 151, 41))
        font = QtGui.QFont()
        font.setBold(True)
        font.setWeight(75)
        self.NOT.setFont(font)
        self.NOT.setStyleSheet("color: rgb(255, 255, 255);\n"
"background-color: rgb(85, 85, 255);")
        self.NOT.setObjectName("NOT")
        self.nhaptudong = QtWidgets.QFrame(self.chonkieunhaplieu)
        self.nhaptudong.setGeometry(QtCore.QRect(-1, 139, 331, 111))
        self.nhaptudong.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.nhaptudong.setFrameShadow(QtWidgets.QFrame.Raised)
        self.nhaptudong.setObjectName("nhaptudong")
        self.YES = QtWidgets.QPushButton(self.nhaptudong)
        self.YES.setGeometry(QtCore.QRect(90, 10, 151, 41))
        font = QtGui.QFont()
        font.setBold(True)
        font.setWeight(75)
        self.YES.setFont(font)
        self.YES.setStyleSheet("color: rgb(255, 255, 255);\n"
"background-color: rgb(85, 85, 255);")
        self.YES.setObjectName("YES")
        self.retranslateUi_YesOrNo(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)
        self.NOT.clicked.connect(self.generate_Phas)
        self.yesOrNoWindow = Form
        
    def __init__(self, selected_text, EXCEL_TABLE):
        self.selected_text = selected_text
        self.EXCEL_TABLE = EXCEL_TABLE
        self.msgBox = None
        self.thread = LongOperationThread(self.selected_text)
        self.thread.finished.connect(self.on_finished)
        self.yesOrNoWindow = None
        
    def func_ResultFinal(self, df):
        self.Form = QtWidgets.QWidget()
        self.ui = Ui_Form_ResultFinal()
        self.ui.setupUi_ResultFinal(self.Form, df)
        self.Form.show()
        self.msgBox.close()
        
    
    def generate_Phas(self, df):
        self.msgBox = QMessageBox()
        self.msgBox.setText("đang tính toán...")
        self.msgBox.setStandardButtons(QMessageBox.NoButton)
        self.msgBox.show()


        self.thread = LongOperationThread(self.selected_text)
        self.thread.finished.connect(self.on_finished)
        self.thread.start()
    

    def on_finished(self, df):
        self.EXCEL_TABLE.setRowCount(len(df))
        self.EXCEL_TABLE.setColumnCount(len(df.columns))
        for i, row in enumerate(df.values):
            for j, value in enumerate(row):
                self.EXCEL_TABLE.setItem(i, j, QtWidgets.QTableWidgetItem(str(value)))
        self.func_ResultFinal(df)
        self.msgBox.done(QMessageBox.Accepted)
        self.yesOrNoWindow.close()
        

    def retranslateUi_YesOrNo(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Form"))
        self.chonkieunhapdulieu.setText(_translate("Form", "Phát hiện chưa có pha"))
        self.NOT.setText(_translate("Form", "Tiếp tục"))
        self.YES.setText(_translate("Form", "Quay lại"))
        
class Ui_Form_ErrorRate(object):
    def setupUi_ErrorRate(self, Form, df):
        Form.setObjectName("Form")
        Form.resize(922, 552)
        self.frame_2 = QtWidgets.QFrame(Form)
        self.frame_2.setGeometry(QtCore.QRect(0, 0, 931, 552))
        self.frame_2.setStyleSheet("background-color: rgb(0, 0, 139);")
        self.frame_2.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame_2.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_2.setObjectName("frame_2")
        self.chonkieunhapdulieu = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu.setGeometry(QtCore.QRect(30, 10, 541, 61))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu.setFont(font)
        self.chonkieunhapdulieu.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu.setObjectName("chonkieunhapdulieu")
        self.Error_Box = QtWidgets.QLineEdit(self.frame_2)
        self.Error_Box.setGeometry(QtCore.QRect(680, 30, 141, 41))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.Error_Box.setFont(font)
        self.Error_Box.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.Error_Box.setObjectName("Error_Box")
        self.nutdonglai_2 = QtWidgets.QPushButton(self.frame_2)
        self.nutdonglai_2.setGeometry(QtCore.QRect(690, 490, 151, 41))
        font = QtGui.QFont()
        font.setPointSize(16)
        font.setBold(True)
        font.setWeight(75)
        self.nutdonglai_2.setFont(font)
        self.nutdonglai_2.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 139);")
        self.nutdonglai_2.setObjectName("nutdonglai_2")
        self.chonkieunhapdulieu_3 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_3.setGeometry(QtCore.QRect(10, 110, 581, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu_3.setFont(font)
        self.chonkieunhapdulieu_3.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_3.setObjectName("chonkieunhapdulieu_3")
        self.MaxLoad_Box = QtWidgets.QLineEdit(self.frame_2)
        self.MaxLoad_Box.setGeometry(QtCore.QRect(680, 160, 141, 41))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.MaxLoad_Box.setFont(font)
        self.MaxLoad_Box.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.MaxLoad_Box.setObjectName("MaxLoad_Box")
        self.chonkieunhapdulieu_4 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_4.setGeometry(QtCore.QRect(170, 150, 261, 61))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu_4.setFont(font)
        self.chonkieunhapdulieu_4.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_4.setObjectName("chonkieunhapdulieu_4")
        self.chonkieunhapdulieu_5 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_5.setGeometry(QtCore.QRect(840, 10, 71, 71))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu_5.setFont(font)
        self.chonkieunhapdulieu_5.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_5.setObjectName("chonkieunhapdulieu_5")
        self.chonkieunhapdulieu_6 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_6.setGeometry(QtCore.QRect(830, 120, 91, 111))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu_6.setFont(font)
        self.chonkieunhapdulieu_6.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_6.setObjectName("chonkieunhapdulieu_6")
        self.chonkieunhapdulieu_2 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_2.setGeometry(QtCore.QRect(10, 60, 581, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(False)
        font.setItalic(True)
        font.setWeight(50)
        self.chonkieunhapdulieu_2.setFont(font)
        self.chonkieunhapdulieu_2.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_2.setObjectName("chonkieunhapdulieu_2")
        self.chonkieunhapdulieu_7 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_7.setGeometry(QtCore.QRect(30, 200, 631, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(False)
        font.setItalic(True)
        font.setWeight(50)
        self.chonkieunhapdulieu_7.setFont(font)
        self.chonkieunhapdulieu_7.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_7.setObjectName("chonkieunhapdulieu_7")
        self.chonkieunhapdulieu_8 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_8.setGeometry(QtCore.QRect(200, 250, 161, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu_8.setFont(font)
        self.chonkieunhapdulieu_8.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_8.setObjectName("chonkieunhapdulieu_8")
        self.Voltage = QtWidgets.QLineEdit(self.frame_2)
        self.Voltage.setGeometry(QtCore.QRect(680, 270, 141, 41))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.Voltage.setFont(font)
        self.Voltage.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.Voltage.setObjectName("Voltage")
        self.chonkieunhapdulieu_9 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_9.setGeometry(QtCore.QRect(10, 300, 621, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(False)
        font.setItalic(True)
        font.setWeight(50)
        self.chonkieunhapdulieu_9.setFont(font)
        self.chonkieunhapdulieu_9.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_9.setObjectName("chonkieunhapdulieu_9")
        self.chonkieunhapdulieu_10 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_10.setGeometry(QtCore.QRect(840, 250, 61, 71))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu_10.setFont(font)
        self.chonkieunhapdulieu_10.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_10.setObjectName("chonkieunhapdulieu_10")
        self.chonkieunhapdulieu_11 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_11.setGeometry(QtCore.QRect(140, 350, 321, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu_11.setFont(font)
        self.chonkieunhapdulieu_11.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_11.setObjectName("chonkieunhapdulieu_11")
        self.chonkieunhapdulieu_12 = QtWidgets.QLabel(self.frame_2)
        self.chonkieunhapdulieu_12.setGeometry(QtCore.QRect(50, 400, 591, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(False)
        font.setItalic(True)
        font.setWeight(50)
        self.chonkieunhapdulieu_12.setFont(font)
        self.chonkieunhapdulieu_12.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu_12.setObjectName("chonkieunhapdulieu_12")
        self.Cos_Phi = QtWidgets.QLineEdit(self.frame_2)
        self.Cos_Phi.setGeometry(QtCore.QRect(680, 360, 141, 41))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.Cos_Phi.setFont(font)
        self.Cos_Phi.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.Cos_Phi.setObjectName("Cos_Phi")
        self.chonkieunhapdulieu_12.raise_()
        self.chonkieunhapdulieu_9.raise_()
										  
        self.chonkieunhapdulieu_2.raise_()
        self.chonkieunhapdulieu_5.raise_()
        self.chonkieunhapdulieu_4.raise_()
        self.chonkieunhapdulieu_3.raise_()
        self.chonkieunhapdulieu.raise_()
        self.Error_Box.raise_()
        self.nutdonglai_2.raise_()
        self.MaxLoad_Box.raise_()
        self.chonkieunhapdulieu_6.raise_()
        self.chonkieunhapdulieu_8.raise_()
        self.Voltage.raise_()
        self.chonkieunhapdulieu_10.raise_()
        self.chonkieunhapdulieu_11.raise_()
        self.Cos_Phi.raise_()
        self.chonkieunhapdulieu_7.raise_()

        

        self.retranslateUi_ErrorRate(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)
        self.Error_Box.textChanged.connect(self.update_max_error_rate)
        self.MaxLoad_Box.textChanged.connect(self.update_max_load_change)
        self.Voltage.textChanged.connect(self.update_voltage)
        self.Cos_Phi.textChanged.connect(self.update_Cos_Phi)
        self.nutdonglai_2.clicked.connect(self.generate_new_phase)
        self.ErrorRateWindow = Form
        
    def __init__(self, selected_text, EXCEL_TABLE):
        self.selected_text = selected_text
        self.EXCEL_TABLE = EXCEL_TABLE
        self.msgBox = None
        self.max_current = 2
        self.max_load_change = 3
        self.voltageset = 220
        self.cosphi = 1
    
        self.ErrorRateWindow = None

    
    def update_max_error_rate(self, text):
        self.max_current = text

    def update_max_load_change(self, text):
        self.max_load_change = text
    
    def update_voltage(self, text):
        self.voltageset = text
    
    def update_Cos_Phi(self, text):
        self.cosphi = text
    
    def func_ResultFinal(self, df_balanced):
        self.msgBox.close()
        self.ErrorRateWindow.close()
        self.Form = QtWidgets.QWidget()
        self.ui = Ui_Form_ResultFinal()
        self.ui.setupUi_ResultFinal(self.Form, df_balanced) 
        self.Form.show()
        
    def generate_new_phase(self, df):
        self.msgBox = QMessageBox()
        self.msgBox.setText("đang tính toán...")
        self.msgBox.setStandardButtons(QMessageBox.NoButton)
        self.msgBox.show()
    
        self.thread = LongOperationThread2(self.selected_text, self, df)
        self.thread.finished.connect(self.on_finished)
        self.thread.start()

    def show_pie_charts(self, current_old, current_new):
        self.chart_window = QtWidgets.QDialog()  
        self.chart_window.setWindowTitle("So sánh pha hiện tại")
        self.chart_window.setWindowFlags(self.chart_window.windowFlags() | QtCore.Qt.WindowMinimizeButtonHint | QtCore.Qt.WindowMaximizeButtonHint)
        layout = QtWidgets.QVBoxLayout(self.chart_window)
        fig = Figure(figsize=(10, 5))
        ax1 = fig.add_subplot(121)
        ax2 = fig.add_subplot(122)
        colors = ['red', 'yellow', 'blue']
        font_size = 40

        current_old = [0 if pd.isna(x) else x for x in current_old]
        current_new = [0 if pd.isna(x) else x for x in current_new]

        def autopct_format(values):
            def my_format(pct):
           
                total = sum([x for x in values if not pd.isna(x)])
                if total == 0:
                    return ''  
                val = (pct * total / 100.0)
                return '{:.3f}A'.format(val) 
            return my_format

        labels = ['Pha A', 'Pha B', 'Pha C']
        # print("current_old:", current_old)
        # print("current_new:", current_new)
        ax1.pie(current_old, labels=labels, autopct=autopct_format(current_old), startangle=90, colors=colors, textprops={'fontsize': font_size})
        ax1.set_title('Trước khi cân bằng', fontsize=font_size)
        ax2.pie(current_new, labels=labels, autopct=autopct_format(current_new), startangle=90, colors=colors, textprops={'fontsize': font_size})
        ax2.set_title('Sau khi cân bằng', fontsize=font_size)
        canvas = FigureCanvas(fig)
        layout.addWidget(canvas)
        self.chart_window.showMinimized()
        self.chart_window.exec_() 
        
    def on_finished(self, df_balanced, changed_df, best_moved_machines_df,
                        current_old_phase_A, current_old_phase_B, current_old_phase_C,
                        current_new_phase_A, current_new_phase_B, current_new_phase_C,
                        max_diff_new_phase_current, max_diff_old_phase_current,
                        PUI_old, PUI_new):
        self.EXCEL_TABLE.setRowCount(len(df_balanced))
        self.EXCEL_TABLE.setColumnCount(len(df_balanced.columns))
        for i, row in enumerate(df_balanced.values):
            for j, value in enumerate(row):
                self.EXCEL_TABLE.setItem(i, j, QtWidgets.QTableWidgetItem(str(value)))

        self.func_ResultFinal(df_balanced) 
        self.msgBox.done(QMessageBox.Accepted)
        diff_old_current = round(max_diff_old_phase_current,3)
        diff_new_current = round(max_diff_new_phase_current,3)
        current_old_A = round(current_old_phase_A,3)
        current_old_B = round(current_old_phase_B,3)
        current_old_C = round(current_old_phase_C,3)
        current_new_A = round(current_new_phase_A,3)
        current_new_B = round(current_new_phase_B,3)
        current_new_C = round(current_new_phase_C,3)
        current_new_C = round(current_new_phase_C,3)
        average_old = round((current_old_A + current_old_B + current_old_C)/3,3)
        average_new = round((current_new_A + current_new_B + current_new_C)/3,3)
        PUI_old = round((diff_old_current/average_old)*100,3)
        PUI_new = round((diff_new_current/average_new)*100,3)
        print(f"Dòng pha A đổi từ {current_old_A}A sang {current_new_A}A ")
        print(f"Dòng pha B đổi từ {current_old_B}A sang {current_new_B}A ")
        print(f"Dòng pha C đổi từ {current_old_C}A sang {current_new_C}A ")
        print(f"Độ lệch dòng lớn nhất trước khi chuyển là: {diff_old_current}A và sau khi chuyển là {diff_new_current} ")
        print(f"PUI trước khi cân bằng = {PUI_old} và sau khi cân bằng ={PUI_new}")
        

        for index, row in changed_df.iterrows():
            print(f"Đề xuất đổi {row['Tên']} từ pha {row['Pha hiện tại']} sang pha {row['Pha đề xuất']}")
        current_old = [current_old_phase_A, current_old_phase_B, current_old_phase_C]
        current_new = [current_new_phase_A, current_new_phase_B, current_new_phase_C]
        
        
        
        print(f"Model AI respond: \n")
        self.show_pie_charts(current_old, current_new)

    def retranslateUi_ErrorRate(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Form"))
        self.chonkieunhapdulieu.setText(_translate("Form", "Sai số dòng tối đa cho phép:"))
        self.nutdonglai_2.setText(_translate("Form", "Bắt đầu"))
        self.chonkieunhapdulieu_3.setText(_translate("Form", "Số lượng tải tối đa được phép"))
        self.chonkieunhapdulieu_4.setText(_translate("Form", " di chuyển:"))
        self.chonkieunhapdulieu_5.setText(_translate("Form", "(A)"))
        self.chonkieunhapdulieu_6.setText(_translate("Form", "(Tải)"))
        self.chonkieunhapdulieu_2.setText(_translate("Form", "(Mặc định là 2A nếu không nhập)"))
        self.chonkieunhapdulieu_7.setText(_translate("Form", "(Mặc định là 3 tải nếu không nhập)"))
        self.chonkieunhapdulieu_8.setText(_translate("Form", "Điện áp:"))
        self.chonkieunhapdulieu_9.setText(_translate("Form", "(Mặc định là 220V nếu không nhập)"))
        self.chonkieunhapdulieu_10.setText(_translate("Form", "(V)"))
        self.chonkieunhapdulieu_11.setText(_translate("Form", "Hệ số công suất:"))
        self.chonkieunhapdulieu_12.setText(_translate("Form", "(Mặc định là 1 nếu không nhập)"))

   
class LongOperationThread2(QThread):
    finished = pyqtSignal(object, object, object, object, object, object, object, object, object, object, object, object, object)  

    def __init__(self, selected_text, ui_form_error_rate, df):
        QThread.__init__(self)
        self.selected_text = selected_text
        self.ui_form_error_rate = ui_form_error_rate
        self.df = df

    def run(self):
        max_load_change = int(self.ui_form_error_rate.max_load_change)
        max_current = float(self.ui_form_error_rate.max_current)
        voltageset = int(self.ui_form_error_rate.voltageset)
        cosphi = float(self.ui_form_error_rate.cosphi)
        time = 24 * 30  
        voltagesetfinal = voltageset / 1000
        
        df_balanced = AI_Func(self.df)  

        column_name = df_balanced.columns[-4]  

       
        old_distribution_total = df_balanced.groupby('Pha hiện tại')[column_name].sum()
        new_distribution_total = df_balanced.groupby('Pha đề xuất')[column_name].sum() 

        old_distribution_total = pd.to_numeric(old_distribution_total, errors='coerce')
        new_distribution_total = pd.to_numeric(new_distribution_total, errors='coerce')
      
        for phase in ['A', 'B', 'C']:
            if phase not in old_distribution_total:
                old_distribution_total[phase] = 0
            if phase not in new_distribution_total:
                new_distribution_total[phase] = 0

     
        current_old_phase_A = old_distribution_total['A'] / (time * voltagesetfinal * cosphi) if (time * voltagesetfinal * cosphi) != 0 else 0
        current_old_phase_B = old_distribution_total['B'] / (time * voltagesetfinal * cosphi) if (time * voltagesetfinal * cosphi) != 0 else 0
        current_old_phase_C = old_distribution_total['C'] / (time * voltagesetfinal * cosphi) if (time * voltagesetfinal * cosphi) != 0 else 0

        current_new_phase_A = new_distribution_total['A'] / (time * voltagesetfinal * cosphi) if (time * voltagesetfinal * cosphi) != 0 else 0
        current_new_phase_B = new_distribution_total['B'] / (time * voltagesetfinal * cosphi) if (time * voltagesetfinal * cosphi) != 0 else 0
        current_new_phase_C = new_distribution_total['C'] / (time * voltagesetfinal * cosphi) if (time * voltagesetfinal * cosphi) != 0 else 0

        max_diff_new_phase_current = max(
            abs(current_new_phase_A - current_new_phase_B),
            abs(current_new_phase_A - current_new_phase_C),
            abs(current_new_phase_B - current_new_phase_C)
        )
        max_diff_old_phase_current = max(
            abs(current_old_phase_A - current_old_phase_B),
            abs(current_old_phase_A - current_old_phase_C),
            abs(current_old_phase_B - current_old_phase_C)
        )
        
        average_old = round((current_old_phase_A + current_old_phase_B + current_old_phase_C) / 3, 3)
        average_new = round((current_new_phase_A + current_new_phase_B + current_new_phase_C) / 3, 3)

       
        PUI_old = round((max_diff_old_phase_current / average_old) * 100, 3)
        PUI_new = round((max_diff_new_phase_current / average_new) * 100, 3)

        changed_df = df_balanced[df_balanced['Pha hiện tại'] != df_balanced['Pha đề xuất']]  
        best_moved_machines_df = changed_df[['Tên', 'Pha hiện tại', 'Pha đề xuất']].copy()  
        best_moved_machines_df.rename(columns={'Pha hiện tại': 'Pha Cũ', 'Pha đề xuất': 'Pha mới'}, inplace=True)

       
        self.finished.emit(
            df_balanced, 
            changed_df, 
            best_moved_machines_df, 
            current_old_phase_A, current_old_phase_B, current_old_phase_C, 
            current_new_phase_A, current_new_phase_B, current_new_phase_C, 
            max_diff_new_phase_current, max_diff_old_phase_current, 
            PUI_old, PUI_new 
        )
        
class Ui_Form_DataTram(object):
    def setupUi_DataTram(self, Form, df):
        Form.setObjectName("Form")
        Form.resize(1920, 1080)
        Form.setStyleSheet("QFrame {\n"
"    background-color: rgb(85, 170, 255);\n"
"    color: rgb(85, 255, 255);\n"
"}")
        self.verticalLayout = QtWidgets.QVBoxLayout(Form)
        self.verticalLayout.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout.setSpacing(0)
        self.verticalLayout.setObjectName("verticalLayout")
        self.frame = QtWidgets.QFrame(Form)
        self.frame.setStyleSheet("background-color: rgb(0, 0, 139);")
        self.frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame.setObjectName("frame")
        self.EDIT_BUTTON = QtWidgets.QPushButton(self.frame)
        self.EDIT_BUTTON.setGeometry(QtCore.QRect(1160, 890, 191, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.EDIT_BUTTON.setFont(font)
        self.EDIT_BUTTON.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 139);")
        self.EDIT_BUTTON.setObjectName("EDIT_BUTTON")
        self.BACK_BUTTON = QtWidgets.QPushButton(self.frame)
        self.BACK_BUTTON.setGeometry(QtCore.QRect(940, 890, 211, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.BACK_BUTTON.setFont(font)
        self.BACK_BUTTON.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 139);")
        self.BACK_BUTTON.setObjectName("BACK_BUTTON")
        self.NAME_OUTPUT = QtWidgets.QLineEdit(self.frame)
        self.NAME_OUTPUT.setGeometry(QtCore.QRect(120, 230, 271, 61))
        font = QtGui.QFont()
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.NAME_OUTPUT.setFont(font)
        self.NAME_OUTPUT.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 0);")
        self.NAME_OUTPUT.setText("")
        self.NAME_OUTPUT.setObjectName("NAME_OUTPUT")
        self.name = QtWidgets.QLabel(self.frame)
        self.name.setGeometry(QtCore.QRect(120, 165, 131, 61))
        font = QtGui.QFont()
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.name.setFont(font)
        self.name.setStyleSheet("color: rgb(255, 255, 255);")
        self.name.setObjectName("name")
        self.tenappviettat = QtWidgets.QLabel(self.frame)
        self.tenappviettat.setGeometry(QtCore.QRect(0, 0, 1920, 111))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(48)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.tenappviettat.setFont(font)
        self.tenappviettat.setStyleSheet("color: rgb(255, 255, 255);")
        self.tenappviettat.setTextFormat(QtCore.Qt.AutoText)
        self.tenappviettat.setAlignment(QtCore.Qt.AlignCenter)
        self.tenappviettat.setObjectName("tenappviettat")
        self.CAN_DAO_PHA = QtWidgets.QPushButton(self.frame)
        self.CAN_DAO_PHA.setGeometry(QtCore.QRect(1600, 195, 211, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.CAN_DAO_PHA.setFont(font)
        self.CAN_DAO_PHA.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 139);")
        self.CAN_DAO_PHA.setObjectName("CAN_DAO_PHA")
									   
        self.EXCEL_TABLE = QtWidgets.QTableWidget(self.frame)
        self.EXCEL_TABLE.setGeometry(QtCore.QRect(400, 190, 1171, 661))
        font = QtGui.QFont()
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.EXCEL_TABLE.setFont(font)
        self.EXCEL_TABLE.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"gridline-color: rgb(0, 0, 0);\n"
"color: rgb(0, 0, 0);\n"
"border-color: rgb(0, 0, 0);")
        self.EXCEL_TABLE.setObjectName("EXCEL_TABLE")
        self.EXCEL_TABLE.setColumnCount(0)
        self.EXCEL_TABLE.setRowCount(0)
        self.tenappviettat_2 = QtWidgets.QLabel(self.frame)
        self.tenappviettat_2.setGeometry(QtCore.QRect(0, 110, 1920, 61))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.tenappviettat_2.setFont(font)
        self.tenappviettat_2.setStyleSheet("color: rgb(255, 255, 255);")
        self.tenappviettat_2.setTextFormat(QtCore.Qt.AutoText)
        self.tenappviettat_2.setAlignment(QtCore.Qt.AlignCenter)
        self.tenappviettat_2.setObjectName("tenappviettat_2")
        self.tenappviettat.raise_()
        self.EDIT_BUTTON.raise_()
        self.BACK_BUTTON.raise_()
        self.NAME_OUTPUT.raise_()
        self.name.raise_()
        self.CAN_DAO_PHA.raise_()
						  
        self.EXCEL_TABLE.raise_()
        self.tenappviettat_2.raise_()
        self.verticalLayout.addWidget(self.frame)

        self.retranslateUi_DataTram(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)
        self.Form = Form
        self.load_data(df)
        self.BACK_BUTTON.clicked.connect(self.go_back)
        self.pick_tram_form = Ui_Form_PickTram()
        self.EDIT_BUTTON.clicked.connect(self.edit_button_clicked)
        if self.selected_text == "Lê Ngọc Hân ":
            df = pd.read_excel('table1.xlsx')
            self.NAME_OUTPUT.setText("Lê Ngọc Hân")
        elif self.selected_text == "Điều kiện xác định":
            df = pd.read_excel('table2.xlsx')
            self.NAME_OUTPUT.setText("Điều kiện xác định")
            
        self.df = df
        self.CAN_DAO_PHA.clicked.connect(self.replace_table3_with_df)
        if 'Pha' in df.columns:
            self.CAN_DAO_PHA.clicked.connect(self.func_ErrorRate)
        else:
            self.CAN_DAO_PHA.clicked.connect(self.func_YesOrNo)
        
            
    def __init__(self, selected_text, MainWindow):
        self.selected_text = selected_text
        self.MainWindow = MainWindow
       
    
    def go_back(self):
        self.Form.hide() 
        self.MainWindow.show() 
        
    def load_data(self, df):
        if self.selected_text == "Lê Ngọc Hân ":
            df = pd.read_excel('table1.xlsx')
        elif self.selected_text == "Điều kiện xác định":
            df = pd.read_excel('table2.xlsx')
        
        self.EXCEL_TABLE.setRowCount(0)
        self.EXCEL_TABLE.setColumnCount(0)
        self.EXCEL_TABLE.setRowCount(df.shape[0])
        self.EXCEL_TABLE.setColumnCount(df.shape[1])
        self.EXCEL_TABLE.setHorizontalHeaderLabels(df.columns.tolist())
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                self.EXCEL_TABLE.setItem(i, j, QtWidgets.QTableWidgetItem(str(value)))
        
    def func_YesOrNo(self, df):
        if self.selected_text == "Lê Ngọc Hân ":
            df = pd.read_excel('table1.xlsx')
        elif self.selected_text == "Điều kiện xác định":
            df = pd.read_excel('table2.xlsx')
    
        self.Form = QtWidgets.QWidget()
        self.ui = Ui_Form_YesOrNo(self.selected_text, self.EXCEL_TABLE)  
        self.ui.setupUi_YesOrNo(self.Form, df)  
        self.Form.show()
            
    def func_ErrorRate(self, df):
        if self.selected_text == "Lê Ngọc Hân ":
            df = pd.read_excel('table1.xlsx')
        elif self.selected_text == "Điều kiện xác định":
            df = pd.read_excel('table2.xlsx')
    
       
        self.Form = QtWidgets.QWidget()
        self.ui = Ui_Form_ErrorRate(self.selected_text, self.EXCEL_TABLE)
        self.ui.setupUi_ErrorRate(self.Form, df)
        self.Form.show()
        
    def replace_table3_with_df(self, df):
        self.df.to_excel('table3.xlsx', index=False)
        
    def edit_button_clicked(self):
        if self.selected_text == "Lê Ngọc Hân ":
            self.pick_tram_form.load_data_doi_can()
        elif self.selected_text == "Điều kiện xác định":
            self.pick_tram_form.load_data_condition()
        
        self.setupUi_DataTram(self.Form, self.df)
        
    def retranslateUi_DataTram(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Form"))
        self.EDIT_BUTTON.setText(_translate("Form", "Chỉnh sửa"))
        self.BACK_BUTTON.setText(_translate("Form", "Quay lại"))
        self.name.setText(_translate("Form", "Trạm:"))
        self.tenappviettat.setText(_translate("Form", "<html><head/><body><p>Phần mềm chuyển tải cân bằng pha</p></body></html>"))
        self.CAN_DAO_PHA.setText(_translate("Form", "Cân đảo pha"))
        self.tenappviettat_2.setText(_translate("Form", "<html><head/><body><p>Tưởng Gia Huy-Trường đại học điện lực</p></body></html>"))

class Ui_Form_Name(object):
    def setupUi_Name(self, Form):
        Form.setObjectName("Form")
        Form.resize(331, 191)
        self.diententram = QtWidgets.QFrame(Form)
        self.diententram.setGeometry(QtCore.QRect(0, 0, 331, 191))
        self.diententram.setMaximumSize(QtCore.QSize(16777215, 16777215))
        self.diententram.setStyleSheet("background-color: rgb(85, 85, 255);")
        self.diententram.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.diententram.setFrameShadow(QtWidgets.QFrame.Raised)
        self.diententram.setObjectName("diententram")
        self.chidan = QtWidgets.QLabel(self.diententram)
        self.chidan.setGeometry(QtCore.QRect(20, 10, 281, 31))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(15)
        font.setBold(True)
        font.setWeight(75)
        self.chidan.setFont(font)
        self.chidan.setStyleSheet("color: rgb(255, 255, 255);")
        self.chidan.setObjectName("chidan")
        self.vitridiententram = QtWidgets.QLineEdit(self.diententram)
        self.vitridiententram.setGeometry(QtCore.QRect(10, 60, 311, 22))
        self.vitridiententram.setStyleSheet("color: rgb(0, 0, 0);\n"
"background-color: rgb(255, 255, 255);")
        self.vitridiententram.setObjectName("vitridiententram")
        self.dong = QtWidgets.QPushButton(self.diententram)
        self.dong.setGeometry(QtCore.QRect(170, 140, 71, 31))
        font = QtGui.QFont()
        font.setBold(True)
        font.setWeight(75)
        self.dong.setFont(font)
        self.dong.setStyleSheet("color: rgb(255, 255, 255);\n"
"background-color: rgb(85, 85, 255);")
        self.dong.setObjectName("dong")
        self.nhapxong = QtWidgets.QPushButton(self.diententram)
        self.nhapxong.setGeometry(QtCore.QRect(260, 140, 61, 31))
        font = QtGui.QFont()
        font.setBold(True)
        font.setWeight(75)
        self.nhapxong.setFont(font)
        self.nhapxong.setStyleSheet("color: rgb(255, 255, 255);\n"
"background-color: rgb(85, 85, 255);")
        self.nhapxong.setObjectName("nhapxong")

        self.retranslateUi_Name(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)
        

    def retranslateUi_Name(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Form"))
        self.chidan.setText(_translate("Form", "Vui lòng điền tên trạm:"))
        self.dong.setText(_translate("Form", "Đóng"))
        self.nhapxong.setText(_translate("Form", "Xong"))


class Ui_Form_forOldNew(object):
    def setupUi_forOldNew(self, Form):
        Form.setObjectName("Form")
        Form.resize(330, 250)
        self.chonkieunhaplieu = QtWidgets.QFrame(Form)
        self.chonkieunhaplieu.setGeometry(QtCore.QRect(0, 0, 330, 250))
        self.chonkieunhaplieu.setStyleSheet("background-color: rgb(85, 85, 255);")
        self.chonkieunhaplieu.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.chonkieunhaplieu.setFrameShadow(QtWidgets.QFrame.Raised)
        self.chonkieunhaplieu.setObjectName("chonkieunhaplieu")
        self.chonkieunhapdulieu = QtWidgets.QLabel(self.chonkieunhaplieu)
        self.chonkieunhapdulieu.setGeometry(QtCore.QRect(90, 0, 251, 41))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(15)
        font.setBold(True)
        font.setWeight(75)
        self.chonkieunhapdulieu.setFont(font)
        self.chonkieunhapdulieu.setStyleSheet("color: rgb(255, 255, 255);")
        self.chonkieunhapdulieu.setObjectName("chonkieunhapdulieu")
        self.nhapthucong = QtWidgets.QPushButton(self.chonkieunhaplieu)
        self.nhapthucong.setGeometry(QtCore.QRect(90, 60, 151, 41))
        font = QtGui.QFont()
        font.setBold(True)
        font.setWeight(75)
        self.nhapthucong.setFont(font)
        self.nhapthucong.setStyleSheet("color: rgb(255, 255, 255);\n"
"background-color: rgb(85, 85, 255);")
        self.nhapthucong.setObjectName("nhapthucong")
        self.nhaptudong = QtWidgets.QFrame(self.chonkieunhaplieu)
        self.nhaptudong.setGeometry(QtCore.QRect(-1, 139, 331, 111))
        self.nhaptudong.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.nhaptudong.setFrameShadow(QtWidgets.QFrame.Raised)
        self.nhaptudong.setObjectName("nhaptudong")
        self.nhapthucong_2 = QtWidgets.QPushButton(self.nhaptudong)
        self.nhapthucong_2.setGeometry(QtCore.QRect(90, -10, 151, 41))
        font = QtGui.QFont()
        font.setBold(True)
        font.setWeight(75)
        self.nhapthucong_2.setFont(font)
        self.nhapthucong_2.setStyleSheet("color: rgb(255, 255, 255);\n"
"background-color: rgb(85, 85, 255);")
        self.nhapthucong_2.setObjectName("nhapthucong_2")
        self.retranslateUi_forOldNew(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)
        self.nhapthucong.clicked.connect(self.func_Name)
        self.nhapthucong_2.clicked.connect(self.func_PickTram)
        self.ForOldNewForm = Form
    
    def func_Name(self):
        self.Form = QtWidgets.QWidget()
        self.ui = Ui_Form_Name()
        self.ui.setupUi_Name(self.Form)
        self.Form.show()
        self.ForOldNewForm.close()
    
    def func_PickTram(self):
        self.Form = QtWidgets.QWidget()
        self.ui = Ui_Form_PickTram()
        self.ui.setupUi_PickTram(self.Form)
        self.Form.show()
        self.ForOldNewForm.close()
        

    def retranslateUi_forOldNew(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Form"))
        self.chonkieunhapdulieu.setText(_translate("Form", "Chọn kiểu nhập:"))
        self.nhapthucong.setText(_translate("Form", "Nhập mới"))
        self.nhapthucong_2.setText(_translate("Form", "Nhập vào trạm có sẵn"))
        

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(1920, 1080)
									 
        MainWindow.setStyleSheet("background-color: rgb(0, 0, 139);\n"
"color: rgb(255, 255, 255);")
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")							   
        self.tieu_de = QtWidgets.QLabel(self.centralwidget)
        self.tieu_de.setGeometry(QtCore.QRect(610, 430, 381, 61))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(28)
        font.setBold(True)
        font.setWeight(75)
        self.tieu_de.setFont(font)
        self.tieu_de.setStyleSheet("color: rgb(255, 255, 255);")
        self.tieu_de.setAlignment(QtCore.Qt.AlignCenter)
        self.tieu_de.setObjectName("tieu_de")
        self.Chon_tram = QtWidgets.QComboBox(self.centralwidget)
        self.Chon_tram.setGeometry(QtCore.QRect(1020, 430, 341, 61))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setWeight(75)
        self.Chon_tram.setFont(font)
        self.Chon_tram.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(0, 0, 139);")
        self.Chon_tram.setObjectName("Chon_tram")
        self.Chon_tram.addItem("")
        self.Chon_tram.addItem("")
        self.Chon_tram.addItem("")
        self.SAMPLE_EXCEL = QtWidgets.QPushButton(self.centralwidget)
        self.SAMPLE_EXCEL.setGeometry(QtCore.QRect(770, 620, 351, 61))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(18)
        font.setBold(True)
        font.setWeight(75)
        self.SAMPLE_EXCEL.setFont(font)
        self.SAMPLE_EXCEL.setStyleSheet("background-color: rgb(255, 255, );\n"
"color: rgb(0, 0, 139);")
        self.SAMPLE_EXCEL.setObjectName("SAMPLE_EXCEL")
        self.SAMPLE_EXCEL.clicked.connect(self.download_sample_excel)
        self.tenappviettat_3 = QtWidgets.QLabel(self.centralwidget)
        self.tenappviettat_3.setGeometry(QtCore.QRect(0, 110, 1920, 61))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(24)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.tenappviettat_3.setFont(font)
        self.tenappviettat_3.setStyleSheet("color: rgb(255, 255, 255);")
        self.tenappviettat_3.setTextFormat(QtCore.Qt.AutoText)
        self.tenappviettat_3.setAlignment(QtCore.Qt.AlignCenter)
        self.tenappviettat_3.setObjectName("tenappviettat_3")
        self.tenappviettat_4 = QtWidgets.QLabel(self.centralwidget)
        self.tenappviettat_4.setGeometry(QtCore.QRect(0, 0, 1920, 111))
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(48)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.tenappviettat_4.setFont(font)
        self.tenappviettat_4.setStyleSheet("color: rgb(255, 255, 255);")
        self.tenappviettat_4.setTextFormat(QtCore.Qt.AutoText)
        self.tenappviettat_4.setAlignment(QtCore.Qt.AlignCenter)
        self.tenappviettat_4.setObjectName("tenappviettat_4")
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        self.Chon_tram.activated[str].connect(self.on_combobox_changed)
        self.pick_tram_form = Ui_Form_PickTram()
        
    def func_forOldNew(self):
        self.Form = QtWidgets.QWidget()
        self.ui = Ui_Form_forOldNew()
        self.ui.setupUi_forOldNew(self.Form)
        self.Form.show()
    
    def func__DataTram(self, df): 
        QtWidgets.QApplication.closeAllWindows()  
        self.Form = QtWidgets.QWidget()
        self.ui = Ui_Form_DataTram(self.selected_text, MainWindow) 
        self.ui.setupUi_DataTram(self.Form, df)  
        self.Form.show()
    
    def func_PickTram(self):
        self.Form = QtWidgets.QWidget()
        self.ui = Ui_Form_PickTram()
        self.ui.setupUi_PickTram(self.Form)
        self.Form.show()

    def open_docx(self, file_path):
            try:
                subprocess.Popen([file_path], shell=True)
            except Exception as e:
                print(f"Error opening file: {e}")
   
    def on_combobox_changed(self, text):
        self.selected_text = text

        if self.selected_text == "Lê Ngọc Hân ":
            df = pd.read_excel('table1.xlsx')
            self.func__DataTram(df) 

        elif self.selected_text == "Điều kiện xác định":
            if os.path.exists("condition.docx"):
                
                self.open_docx("condition.docx")
            else:
                self.pick_tram_form.load_data_condition()
                
        elif text == "Nhập dữ liệu":
            self.pick_tram_form.load_data('table1')

        # elif text == "Nhập dữ liệu":
        #     self.func_PickTram()
    
    def download_sample_excel(self):
        base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    
        excel_file_path = os.path.join(base_path, 'Sample.xlsx')
    
        save_location, _ = QtWidgets.QFileDialog.getSaveFileName(MainWindow, 'Save File', '', 'Excel Files (*.xlsx)')
    
        shutil.copy(excel_file_path, save_location)


    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.tieu_de.setText(_translate("MainWindow", "Lựa chọn trạm:"))
        self.Chon_tram.setItemText(0, _translate("MainWindow", "Lê Ngọc Hân "))
        self.Chon_tram.setItemText(1, _translate("MainWindow", "Điều kiện xác định"))
        self.Chon_tram.setItemText(2, _translate("MainWindow", "Nhập dữ liệu"))
        self.SAMPLE_EXCEL.setText(_translate("MainWindow", "Download file excel mẫu"))
        self.tenappviettat_3.setText(_translate("MainWindow", "<html><head/><body><p>Tưởng Gia Huy-Trường đại học điện lực</p></body></html>"))                                                                                                                                                
        self.tenappviettat_4.setText(_translate("MainWindow", "<html><head/><body><p>PHẦN MỀM CHUYỂN TẢI CÂN BẰNG PHA</p></body></html>"))

                                                      
                                                                                                                                                                                                                                                                                                                                                   #Credit: Tuong Gia Huy, Nguyen Trung Hieu, Tong Vinh Lap
if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())